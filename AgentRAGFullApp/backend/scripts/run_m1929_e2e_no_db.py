"""Sprint M19.29 · E2E real renderer Claude + docx-js — sin Supabase.

Pipeline completo en LOCAL sin DB:
  1. Lee 3 SKILL.md de TEST Freddy/templates_review/
  2. Por cada uno: construye RenderRequest, llama Claude API real (Opus 4.7)
     → Claude genera JS → sandbox Node ejecuta → .docx real
  3. Guarda .docx en outputs/, mide tokens/costo/duración/retries
  4. Abre cada .docx con python-docx y extrae texto para análisis cualitativo

Salida:
  - outputs/<doc_type>_<ts>.docx (los .docx reales)
  - outputs/<doc_type>_<ts>.js (el JS que generó Claude, para auditar)
  - outputs/REPORTE_E2E_M1929.md (análisis humano-legible)

Uso:
    set ANTHROPIC_API_KEY=<key>
    set LEXAI_NODE_PATH_OVERRIDE=c:/tmp/docx-runtime/node_modules
    python scripts/run_m1929_e2e_no_db.py
"""

from __future__ import annotations

import asyncio
import io
import os
import sys
import time
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))


# ---------------------------------------------------------------------------
# Templates a probar (3 doc_types diferentes para comparativa)
# ---------------------------------------------------------------------------

TEMPLATES_DIR = Path(os.getenv("LEXAI_TEMPLATES_REVIEW_DIR") or
    "C:/Users/freddyrs/Desktop/Legal Demo/Legal_agent_Frontend/TEST Freddy/templates_review")

OUTPUTS_DIR = Path(
    "C:/Users/freddyrs/Desktop/Legal Demo/Legal_agent_Frontend/TEST Freddy/test_outputs"
)

# 3 escenarios — variedad de complejidad
SCENARIOS: list[dict[str, Any]] = [
    {
        "label": "Notarial · Revocatoria de Poder",
        "template_file": "03_notarial_revocatoria_poder_co.md",
        "doc_type": "revocatoria_poder",
        "user_prompt": (
            "Redacta una revocatoria de poder en Bogotá, fechada 28 de mayo de 2026. "
            "El poderdante es Juan Pérez García (CC 79.123.456) y revoca el poder "
            "otorgado el 10 de marzo de 2024 ante la Notaría 19 de Bogotá (Escritura 1234) "
            "a María López Ruiz (CC 52.987.654). El poder revocado le permitía representarme "
            "en procesos civiles ante juzgados de Bogotá."
        ),
        "data": {
            "ciudad": "Bogotá",
            "fecha": "28 de mayo de 2026",
            "nombre_poderdante": "Juan Pérez García",
            "cc_poderdante": "79.123.456",
            "nombre_apoderado": "María López Ruiz",
            "cc_apoderado": "52.987.654",
            "fecha_poder_original": "10 de marzo de 2024",
            "notaria_original": "Notaría 19 de Bogotá",
            "numero_escritura_original": "1234",
            "sintesis_facultades": "representarme en procesos civiles ante juzgados de Bogotá",
        },
    },
    {
        "label": "Contractual · Compraventa de Vehículo",
        "template_file": "18_contractual_compraventa_vehiculo_co.md",
        "doc_type": "compraventa_vehiculo",
        "user_prompt": (
            "Redacta contrato de compraventa de vehículo. Vendedor: Carlos Gómez (CC 70.111.222). "
            "Comprador: Ana Torres (CC 41.555.666). Vehículo: Toyota Hilux 2022, blanco, "
            "placa ABC123, kilometraje 45000. Precio $95.000.000 pagaderos de contado el día de "
            "la entrega (15 de junio de 2026 en Medellín)."
        ),
        "data": {
            "fecha": "28 de mayo de 2026",
            "ciudad": "Medellín",
            "nombre_vendedor": "Carlos Gómez",
            "cc_vendedor": "70.111.222",
            "nombre_comprador": "Ana Torres",
            "cc_comprador": "41.555.666",
            "placa": "ABC123",
            "marca": "Toyota",
            "modelo": "2022",
            "linea": "Hilux",
            "color": "Blanco",
            "numero_motor": "1GDFTV12345",
            "numero_chasis": "MR0FR22G500123456",
            "cilindraje": "2700",
            "tipo_carroceria": "Pickup",
            "kilometraje": "45000",
            "precio_numero": "$95.000.000",
            "precio_letras": "noventa y cinco millones de pesos",
            "fecha_entrega": "15 de junio de 2026",
        },
    },
    {
        "label": "Petitorio · Requerimiento Extrajudicial",
        "template_file": "15_petitorio_requerimiento_extrajudicial_co.md",
        "doc_type": "requerimiento_extrajudicial",
        "user_prompt": (
            "Redacta requerimiento extrajudicial de cobro. Acreedor: Comercializadora Andina "
            "SAS (NIT 900.111.222-3). Deudor: Industrias del Norte SAS (NIT 800.444.555-1). "
            "Origen: factura 7821 del 15 de febrero de 2026 por $25.500.000. Vencida hace "
            "100 días. Solicitar pago en 10 días hábiles."
        ),
        "data": {
            "fecha": "28 de mayo de 2026",
            "ciudad": "Bogotá",
            "nombre_acreedor": "Comercializadora Andina SAS",
            "cc_nit_acreedor": "900.111.222-3",
            "nombre_deudor": "Industrias del Norte SAS",
            "cc_nit_deudor": "800.444.555-1",
            "tipo_documento_obligacion": "factura",
            "numero_obligacion": "7821",
            "fecha_obligacion": "15 de febrero de 2026",
            "capital_adeudado": "$25.500.000",
            "capital_letras": "veinticinco millones quinientos mil pesos",
            "fecha_mora": "16 de marzo de 2026",
            "plazo_pago_dias": "10",
            "cuenta_bancaria": "Banco de Bogotá cuenta corriente 123-456789-01",
        },
    },
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _now_ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _read_template(filename: str) -> tuple[dict[str, Any], str]:
    """Lee un SKILL.md de TEMPLATES_DIR y devuelve (frontmatter, body completo)."""
    path = TEMPLATES_DIR / filename
    md = path.read_text(encoding="utf-8")
    # Parser mínimo (idem scripts/seed_12_new_templates_co.py)
    import re
    m = re.match(r"^---\s*\n(.*?)\n---\s*\n(.*)$", md.strip(), re.DOTALL)
    fm: dict[str, Any] = {}
    body = md
    if m:
        raw_yaml, body = m.group(1), m.group(2).strip()
        for line in raw_yaml.splitlines():
            if ":" in line and not line.lstrip().startswith("#"):
                k, _, v = line.partition(":")
                fm[k.strip()] = v.strip().strip('"').strip("'")
    return fm, body


def _extract_docx_text(docx_bytes: bytes) -> str:
    """Extrae texto plano del .docx para análisis cualitativo."""
    try:
        from docx import Document  # python-docx
        doc = Document(io.BytesIO(docx_bytes))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"<error extracting docx: {e}>"


def _analyze_quality(scenario: dict[str, Any], extracted_text: str) -> dict[str, Any]:
    """Heurísticas simples de calidad."""
    data = scenario["data"]
    # Cuántos placeholders esperados aparecen
    found_keys = 0
    total_keys = 0
    found_values: list[str] = []
    missing_values: list[str] = []
    for k, v in data.items():
        if isinstance(v, str) and len(v) > 3:
            total_keys += 1
            if v in extracted_text:
                found_keys += 1
                found_values.append(k)
            else:
                missing_values.append(f"{k}={v[:30]}")

    # Placeholders sin sustituir
    import re
    unfilled = re.findall(r"\[[A-Z_]{3,40}\]", extracted_text)

    # Líneas y palabras
    lines = extracted_text.count("\n") + 1
    words = len(extracted_text.split())

    # Detectar emojis (prohibidos)
    emoji_chars = re.findall(r"[\U0001F300-\U0001F9FF☀-➿]", extracted_text)

    return {
        "data_substitution_rate": f"{found_keys}/{total_keys} = {(100*found_keys/total_keys if total_keys else 0):.0f}%",
        "found_values": found_values,
        "missing_values": missing_values,
        "unfilled_placeholders": unfilled[:10],
        "unfilled_count": len(unfilled),
        "char_count": len(extracted_text),
        "word_count": words,
        "line_count": lines,
        "emoji_count": len(emoji_chars),
    }


def _validate_docx(docx_bytes: bytes) -> dict[str, Any]:
    """Validaciones técnicas del .docx (zip, parts, ...)."""
    out: dict[str, Any] = {
        "size_bytes": len(docx_bytes),
        "magic_ok": docx_bytes.startswith(b"PK\x03\x04"),
    }
    try:
        with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
            names = zf.namelist()
            out["zip_entries"] = len(names)
            out["has_document_xml"] = "word/document.xml" in names
            out["has_styles_xml"] = any("word/styles.xml" in n for n in names)
            out["has_content_types"] = "[Content_Types].xml" in names
            doc_xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
            out["document_xml_bytes"] = len(doc_xml)
    except Exception as e:
        out["zip_error"] = str(e)
    return out


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


async def run_scenario(scenario: dict[str, Any]) -> dict[str, Any]:
    from lex.renderer.claude_docx_renderer import RenderRequest, render_docx
    from lex.renderer.node_executor import NodeExecutionError

    label = scenario["label"]
    print(f"\n{'='*72}\n[SCENARIO] {label}\n{'='*72}")

    fm, body = _read_template(scenario["template_file"])
    template_md = body
    template_chars = len(template_md)
    print(f"  template_chars={template_chars}  (approx tokens {template_chars//4})")

    req = RenderRequest(
        doc_type=scenario["doc_type"],
        template_skill_md=template_md,
        user_prompt=scenario["user_prompt"],
        data=scenario["data"],
        firm_id=None,
        user_id=None,
        max_retries=2,
        timeout_s=60,
        temperature=0.15,
    )

    started = time.perf_counter()
    error_info: dict[str, Any] = {}
    result = None
    try:
        result = await render_docx(req)
    except NodeExecutionError as e:
        error_info = {"kind": "sandbox", "node_kind": e.kind, "msg": str(e), "stderr": (e.stderr or "")[:600]}
    except Exception as e:
        error_info = {"kind": "pipeline", "exc": type(e).__name__, "msg": str(e)[:500]}
    elapsed = int((time.perf_counter() - started) * 1000)

    out: dict[str, Any] = {
        "label": label,
        "doc_type": scenario["doc_type"],
        "template_chars": template_chars,
        "user_prompt_chars": len(scenario["user_prompt"]),
        "total_duration_ms": elapsed,
        "error": error_info or None,
    }

    if not result:
        print(f"  FAILED: {error_info}")
        return out

    # Persistir .docx y .js
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    ts = _now_ts()
    docx_path = OUTPUTS_DIR / f"{scenario['doc_type']}_{ts}.docx"
    js_path = OUTPUTS_DIR / f"{scenario['doc_type']}_{ts}.js"
    docx_path.write_bytes(result.docx_bytes)
    js_path.write_text(result.generated_js, encoding="utf-8")

    # Validar + analizar
    tech = _validate_docx(result.docx_bytes)
    extracted = _extract_docx_text(result.docx_bytes)
    quality = _analyze_quality(scenario, extracted)

    out.update({
        "docx_path": str(docx_path),
        "js_path": str(js_path),
        "docx_sha256": result.sha256,
        "js_sha256": result.js_sha256,
        "js_bytes": len(result.generated_js.encode("utf-8")),
        "retry_count": result.retry_count,
        "llm_duration_ms": result.llm_duration_ms,
        "sandbox_duration_ms": result.sandbox_duration_ms,
        "tokens_input": result.tokens_input,
        "tokens_output": result.tokens_output,
        "cost_usd_cents": result.cost_usd_cents,
        "model": result.model,
        "tech": tech,
        "quality": quality,
        "extracted_text_excerpt": extracted[:600],
    })
    print(f"  OK · docx={tech.get('size_bytes')}B retries={result.retry_count} "
          f"llm_ms={result.llm_duration_ms} sandbox_ms={result.sandbox_duration_ms} "
          f"tokens_in={result.tokens_input} tokens_out={result.tokens_output} "
          f"cost_cents={result.cost_usd_cents}")
    print(f"  data_subst={quality['data_substitution_rate']} "
          f"unfilled={quality['unfilled_count']} chars={quality['char_count']}")
    return out


def write_report(results: list[dict[str, Any]]) -> Path:
    rpath = OUTPUTS_DIR / f"REPORTE_E2E_M1929_{_now_ts()}.md"
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)

    success = [r for r in results if not r.get("error")]
    failed = [r for r in results if r.get("error")]
    total_cents = sum(r.get("cost_usd_cents", 0) for r in success)
    total_in = sum(r.get("tokens_input", 0) for r in success)
    total_out = sum(r.get("tokens_output", 0) for r in success)

    lines: list[str] = [
        "# Reporte E2E M19.29 · Renderer Claude + docx-js (camino B)",
        "",
        f"Fecha: {datetime.now().isoformat(timespec='seconds')}  ",
        f"Modelo: {(success[0] if success else {}).get('model', 'N/A')}  ",
        "",
        "## Resumen ejecutivo",
        "",
        f"- Escenarios ejecutados: **{len(results)}**",
        f"- Éxitos: **{len(success)}** · Fallos: **{len(failed)}**",
        f"- Tokens totales — input: **{total_in}** · output: **{total_out}**",
        f"- Costo total estimado: **${total_cents/100:.4f} USD** ({total_cents} cents)",
        f"- Duración promedio LLM: **{sum(r.get('llm_duration_ms',0) for r in success)//max(1,len(success))} ms**",
        f"- Duración promedio sandbox: **{sum(r.get('sandbox_duration_ms',0) for r in success)//max(1,len(success))} ms**",
        "",
        "## Configuración",
        "",
        f"- Templates dir: `{TEMPLATES_DIR}`",
        f"- Outputs dir: `{OUTPUTS_DIR}`",
        f"- Node sandbox: NODE_PATH = `{os.environ.get('NODE_PATH','(default)')}`",
        f"- Feature flag CLAUDE_RENDERER_ENABLED: `{os.environ.get('CLAUDE_RENDERER_ENABLED','(unset)')}`",
        "",
    ]

    for i, r in enumerate(results, start=1):
        lines += [
            f"## Escenario {i}: {r['label']}",
            "",
            f"- doc_type: `{r['doc_type']}`",
            f"- template_chars: {r['template_chars']} (~{r['template_chars']//4} tokens)",
            f"- user_prompt_chars: {r['user_prompt_chars']}",
            f"- total_duration_ms: {r['total_duration_ms']}",
            "",
        ]
        if r.get("error"):
            lines += [
                "**Resultado: FAIL**",
                "",
                f"```\n{r['error']}\n```",
                "",
            ]
            continue
        tech = r["tech"]
        q = r["quality"]
        lines += [
            "**Resultado: OK**",
            "",
            f"- .docx → `{r['docx_path']}` ({tech['size_bytes']} bytes, sha256 `{r['docx_sha256'][:16]}…`)",
            f"- .js generado → `{r['js_path']}` ({r['js_bytes']} bytes)",
            f"- retries: {r['retry_count']} · llm_ms: {r['llm_duration_ms']} · sandbox_ms: {r['sandbox_duration_ms']}",
            f"- tokens_in: {r['tokens_input']} · tokens_out: {r['tokens_output']} · cost_cents: {r['cost_usd_cents']}",
            "",
            "### Validación técnica del .docx",
            "",
            f"- magic_bytes_OK: {tech['magic_ok']}",
            f"- zip_entries: {tech.get('zip_entries')}",
            f"- has_document_xml: {tech.get('has_document_xml')}",
            f"- has_styles_xml: {tech.get('has_styles_xml')}",
            f"- document_xml_bytes: {tech.get('document_xml_bytes')}",
            "",
            "### Análisis cualitativo",
            "",
            f"- Sustitución de datos del usuario: **{q['data_substitution_rate']}**",
            f"- Datos NO encontrados ({len(q['missing_values'])}): {', '.join(q['missing_values'][:6]) if q['missing_values'] else 'ninguno'}",
            f"- Placeholders `[XXX]` sin sustituir: **{q['unfilled_count']}** {q['unfilled_placeholders'][:6]}",
            f"- Tamaño documento: {q['char_count']} chars · {q['word_count']} palabras · {q['line_count']} líneas",
            f"- Emojis detectados (deben ser 0): **{q['emoji_count']}**",
            "",
            "### Excerpt del texto generado (primeros 600 chars)",
            "",
            "```",
            r["extracted_text_excerpt"],
            "```",
            "",
        ]

    lines += [
        "## Conclusiones",
        "",
        f"El pipeline camino B (Claude + docx-js@9.5 en sandbox Node 20) generó "
        f"{len(success)}/{len(results)} documentos .docx válidos.",
        "",
    ]
    if failed:
        lines += [f"- {len(failed)} escenario(s) fallaron — revisar `error` arriba para root cause."]
    if success:
        avg_subst = sum(
            int(r["quality"]["data_substitution_rate"].split("=")[1].strip().rstrip("%"))
            for r in success
        ) / len(success)
        lines += [f"- Tasa promedio de sustitución de datos: **{avg_subst:.0f}%**"]
        avg_unfilled = sum(r["quality"]["unfilled_count"] for r in success) / len(success)
        lines += [f"- Promedio placeholders sin sustituir por doc: **{avg_unfilled:.1f}**"]

    rpath.write_text("\n".join(lines), encoding="utf-8")
    print(f"\n[REPORT] {rpath}")
    return rpath


async def main() -> int:
    if not os.environ.get("ANTHROPIC_API_KEY", "").strip():
        print("ERROR: ANTHROPIC_API_KEY no está seteada", file=sys.stderr)
        return 2
    if not TEMPLATES_DIR.exists():
        print(f"ERROR: TEMPLATES_DIR missing: {TEMPLATES_DIR}", file=sys.stderr)
        return 2

    os.environ.setdefault("CLAUDE_RENDERER_ENABLED", "true")
    np = os.environ.get("LEXAI_NODE_PATH_OVERRIDE", "").strip()
    if np:
        os.environ["NODE_PATH"] = np

    results: list[dict[str, Any]] = []
    for sc in SCENARIOS:
        r = await run_scenario(sc)
        results.append(r)
    write_report(results)
    return 0 if all(not r.get("error") for r in results) else 1


if __name__ == "__main__":
    sys.exit(asyncio.run(main()))
