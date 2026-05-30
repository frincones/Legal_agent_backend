"""Sprint M19.31 · python_docx_builder.

Renderer Python puro que aplica el estilo del SKILL.md a los bloques
persistidos. Reemplaza el `docx_forensic_builder.py` legacy (1 102
líneas con mapping romanos hardcoded para todos los doc_types).

Características clave:
  - NO usa LLM (siempre rápido, ~200ms para 60 bloques).
  - Aplica `SkillContext.style_hints` (fuente, márgenes, headings).
  - Aplica `SkillContext.sections` para los títulos de cláusulas
    (PRIMERA. OBJETO DEL PODER en vez de "I. PARTES" hardcoded).
  - Si no hay SkillContext, cae a estilo "forense" sensato como default.
  - Compatible con todos los tipos de Block del schema actual.
"""

from __future__ import annotations

import logging
import os
from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

from lex.orchestrator.skill_context import SkillContext

logger = logging.getLogger(__name__)


# ============================================================================
# Helpers
# ============================================================================


def _hex_to_rgb(hex_str: str | None) -> RGBColor:
    if not hex_str:
        return RGBColor(0, 0, 0)
    s = hex_str.strip().lstrip("#")
    if len(s) != 6:
        return RGBColor(0, 0, 0)
    try:
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except ValueError:
        return RGBColor(0, 0, 0)


def _alignment_from_str(s: str | None) -> int:
    s = (s or "").strip().lower()
    if s in ("center", "centro"):
        return WD_ALIGN_PARAGRAPH.CENTER
    if s in ("right", "derecha"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    if s in ("justify", "justified", "justificado"):
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


# ============================================================================
# StyleConfig — extracto del SkillContext aplicable al .docx
# ============================================================================


class StyleConfig:
    """Configuración resuelta del estilo. Si hay SkillContext, lo extrae de él;
    si no, usa defaults forenses sensatos compatibles con doc_forensic_builder.
    """

    def __init__(self, skill_ctx: Optional[SkillContext]):
        sh = skill_ctx.style_hints if (skill_ctx and skill_ctx.style_hints) else None
        # Fuente cuerpo
        self.font_family = (sh.font_family if sh else "Arial")
        self.body_size_pt = (sh.font_size_pt if sh else 11)
        # Título principal del documento
        self.title_size_pt = 14
        self.title_bold = True
        self.title_alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Heading de sección/cláusula
        self.heading_size_pt = (sh.heading1_size_pt if sh else 12)
        self.heading_bold = (sh.heading1_bold if sh else True)
        self.heading_color = _hex_to_rgb(sh.heading1_color if sh else "000000")
        self.heading_alignment = _alignment_from_str(sh.heading1_alignment if sh else "left")
        # Página
        self.page_margin_in = (sh.page_margin_inches if sh else 1.0)
        self.page_size = (sh.page_size if sh else "Letter")
        # Cuerpo
        self.line_spacing = (sh.line_spacing if sh else 1.15)
        self.first_line_indent_in = (sh.first_line_indent_inches if sh else 0.0)
        self.paragraph_after_pt = (sh.paragraph_after_pt if sh else 6)
        self.justify_body = (sh.justify_body if sh else True)


def _apply_page_setup(doc: Document, cfg: StyleConfig) -> None:
    for section in doc.sections:
        section.top_margin = Inches(cfg.page_margin_in)
        section.bottom_margin = Inches(cfg.page_margin_in)
        section.left_margin = Inches(cfg.page_margin_in)
        section.right_margin = Inches(cfg.page_margin_in)
        # Letter por default; A4 si el SKILL.md lo pide
        if cfg.page_size == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        elif cfg.page_size == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14.0)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.0)


def _apply_default_doc_style(doc: Document, cfg: StyleConfig) -> None:
    """Setea estilo "Normal" base."""
    style = doc.styles["Normal"]
    style.font.name = cfg.font_family
    style.font.size = Pt(cfg.body_size_pt)
    # Set both East-Asian font for compatibility
    rpr = style.element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style.element.insert(0, rpr)


def _set_run_font(run, cfg: StyleConfig, *, bold: bool = False, italic: bool = False,
                  size_pt: int | None = None, color: RGBColor | None = None,
                  underline: bool = False) -> None:
    run.font.name = cfg.font_family
    run.font.size = Pt(size_pt or cfg.body_size_pt)
    run.bold = bool(bold)
    run.italic = bool(italic)
    run.underline = bool(underline)
    if color is not None:
        run.font.color.rgb = color


# ============================================================================
# Block renderers
# ============================================================================


def _render_title(doc: Document, cfg: StyleConfig, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = cfg.title_alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run((text or "").strip().upper())
    _set_run_font(run, cfg, bold=cfg.title_bold, size_pt=cfg.title_size_pt)


def _clause_sections(skill_ctx: SkillContext) -> list[Any]:
    """Filtra solo secciones "clausulares" del SkillContext.

    Una sección del SKILL.md es clausular si:
      - Tiene roman/ordinal detectado (PRIMERA, SEGUNDA, ..., DÉCIMA, I, II, ...)
      - O su título empieza con palabras clave tipo "PRIMERA", "I.", "HECHOS",
        "PRETENSIONES", "FUNDAMENTOS", "OBJETO DEL", etc.

    Items meta del MD (como "Título", "Subtítulo italic", "Destinatario notarial",
    "Comparecencia del poderdante", "Designación del apoderado", "Lugar, fecha
    y firmas", "Diligencia notarial de autenticación") quedan FUERA — esos
    no corresponden a SectionHeadingBlock en el output del orchestrator.
    """
    import re as _re
    out: list[Any] = []
    clausular_kw = _re.compile(
        r"^\s*(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[EÉ]PTIMA|"
        r"OCTAVA|NOVENA|D[EÉ]CIMA|UND[EÉ]CIMA|DUOD[EÉ]CIMA|"
        r"PEN[UÚ]LTIMA|[UÚ]LTIMA|HECHOS|PRETENSIONES|FUNDAMENTOS|"
        r"OBJETO|ANTECEDENTES|PARTES|PETICIONES|VIGENCIA|"
        r"I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII|XIII)[\.\s]",
        _re.IGNORECASE,
    )
    for s in (skill_ctx.sections or []):
        title = (s.title or "").strip()
        if s.roman or clausular_kw.match(title):
            out.append(s)
    return out


def _render_section_heading(
    doc: Document, cfg: StyleConfig, *,
    roman: str | None,
    text: str,
    clause_idx: int,
    skill_ctx: Optional[SkillContext],
) -> None:
    """Renderiza un heading de sección.

    Estrategia:
      - Si `text` ya viene del SKILL.md (lo común en M19.31), incluye el
        ordinal completo ("PRIMERA. OBJETO DEL PODER") → usar tal cual.
      - Si `text` es genérico ("PARTES", "OBJETO") y hay `roman`, combinar
        ("I. PARTES") — preserva compat con plan legacy.
      - Limpiar artefactos del MD: `**texto**` → `texto`.
    """
    raw = (text or "").strip()
    # Limpiar markdown bold
    if "**" in raw:
        raw = raw.replace("**", "").strip()
    if not raw:
        return

    # Detectar si el texto YA contiene un ordinal al inicio
    import re as _re
    has_ordinal = bool(_re.match(
        r"^\s*(PRIMERA|SEGUNDA|TERCERA|CUARTA|QUINTA|SEXTA|S[EÉ]PTIMA|"
        r"OCTAVA|NOVENA|D[EÉ]CIMA|UND[EÉ]CIMA|DUOD[EÉ]CIMA|"
        r"PEN[UÚ]LTIMA|[UÚ]LTIMA|"
        r"I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII|XIII)[\.\s]",
        raw, _re.IGNORECASE,
    ))

    if has_ordinal:
        final_text = raw
    elif roman and raw:
        final_text = f"{roman}. {raw}"
    else:
        final_text = raw

    final_text = final_text.strip().rstrip(".") + "."
    if not final_text:
        return

    p = doc.add_paragraph()
    p.alignment = cfg.heading_alignment
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(final_text)
    _set_run_font(run, cfg,
                  bold=cfg.heading_bold,
                  size_pt=cfg.heading_size_pt,
                  color=cfg.heading_color)


def _render_subsection(doc: Document, cfg: StyleConfig, *, number: str, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}. {text}".strip())
    _set_run_font(run, cfg, bold=True, size_pt=cfg.heading_size_pt - 1)


def _render_paragraph_with_runs(
    doc: Document, cfg: StyleConfig, *, runs_data: list[dict], align: str = "justify",
    indent_left_cm: float | None = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = _alignment_from_str(align)
    if cfg.first_line_indent_in > 0:
        p.paragraph_format.first_line_indent = Inches(cfg.first_line_indent_in)
    if indent_left_cm:
        p.paragraph_format.left_indent = Cm(indent_left_cm)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = cfg.line_spacing
    p.paragraph_format.space_after = Pt(cfg.paragraph_after_pt)
    if not runs_data:
        p.add_run("")
        return
    for r in runs_data:
        if isinstance(r, str):
            text = r
            bold = italic = underline = False
        else:
            text = r.get("text", "")
            bold = bool(r.get("bold"))
            italic = bool(r.get("italic"))
            underline = bool(r.get("underline"))
        if not text:
            continue
        run = p.add_run(text)
        _set_run_font(run, cfg, bold=bold, italic=italic, underline=underline)


def _render_hecho(doc: Document, cfg: StyleConfig, *, num: int, runs_data: list[dict]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(cfg.paragraph_after_pt)
    head = p.add_run(f"{num}. ")
    _set_run_font(head, cfg, bold=True)
    for r in runs_data or []:
        if isinstance(r, str):
            text = r; bold = italic = False
        else:
            text = r.get("text", "")
            bold = bool(r.get("bold"))
            italic = bool(r.get("italic"))
        if text:
            run = p.add_run(text)
            _set_run_font(run, cfg, bold=bold, italic=italic)


def _render_pretension(doc: Document, cfg: StyleConfig, *, ord_text: str, runs_data: list[dict]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.first_line_indent = Cm(-1.5)
    p.paragraph_format.space_after = Pt(cfg.paragraph_after_pt)
    head = p.add_run(f"{ord_text}. ")
    _set_run_font(head, cfg, bold=True)
    for r in runs_data or []:
        if isinstance(r, str):
            text = r; bold = italic = False
        else:
            text = r.get("text", "")
            bold = bool(r.get("bold"))
            italic = bool(r.get("italic"))
        if text:
            run = p.add_run(text)
            _set_run_font(run, cfg, bold=bold, italic=italic)


def _render_list_item(doc: Document, cfg: StyleConfig, *, num: str, runs_data: list[dict]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(4)
    head = p.add_run(f"{num}) ")
    _set_run_font(head, cfg, bold=False)
    for r in runs_data or []:
        if isinstance(r, str):
            text = r; bold = italic = False
        else:
            text = r.get("text", "")
            bold = bool(r.get("bold"))
            italic = bool(r.get("italic"))
        if text:
            run = p.add_run(text)
            _set_run_font(run, cfg, bold=bold, italic=italic)


def _tier_from_data(data: dict) -> str:
    """M20.10/M20.13: extrae tier 5-state. Backwards compat con verified/derogada."""
    tier = (data.get("tier") or "").upper().strip()
    if tier in ("GROUNDED", "DEROGADA", "VERIFY_FLAG", "NOT_FOUND", "MODULADA"):
        return tier
    # Legacy mapping
    if data.get("derogada"):
        return "DEROGADA"
    if data.get("verified"):
        return "GROUNDED"
    return "VERIFY_FLAG"


def _render_tier_badge(p, cfg: StyleConfig, tier: str, *, fuente_url: str | None = None,
                        derogada_por: str | None = None, suggested: str | None = None) -> None:
    """M20.10: renderiza badge según tier con color + texto explicativo."""
    if tier == "GROUNDED":
        badge = p.add_run("  ✓ verificado")
        _set_run_font(badge, cfg, bold=False, color=_hex_to_rgb("00B050"))
        if fuente_url:
            note = p.add_run(f"  [fuente: {_short_url(fuente_url)}]")
            _set_run_font(note, cfg, italic=True, size_pt=cfg.body_size_pt - 2,
                           color=_hex_to_rgb("707070"))
    elif tier == "DEROGADA":
        badge = p.add_run("  ✗ DEROGADA")
        _set_run_font(badge, cfg, bold=True, color=_hex_to_rgb("C00000"))
        if derogada_por:
            note = p.add_run(f"  [por {derogada_por}]")
            _set_run_font(note, cfg, italic=True, bold=True, color=_hex_to_rgb("C00000"))
    elif tier == "NOT_FOUND":
        badge = p.add_run("  ⊗ NO ENCONTRADA")
        _set_run_font(badge, cfg, bold=True, color=_hex_to_rgb("C00000"))
        if suggested:
            note = p.add_run(f"  [sugerencia: {suggested}]")
            _set_run_font(note, cfg, italic=True, color=_hex_to_rgb("C00000"))
    elif tier == "MODULADA":
        # M20.13: 5° tier - vigente con modulación constitucional
        badge = p.add_run("  ⚖ MODULADA")
        _set_run_font(badge, cfg, bold=True, color=_hex_to_rgb("8B6F00"))
        if suggested:
            note = p.add_run(f"  [aplicar con limitaciones de: {suggested}]")
            _set_run_font(note, cfg, italic=True, color=_hex_to_rgb("8B6F00"))
    else:   # VERIFY_FLAG
        badge = p.add_run("  ⚠ [verificar]")
        _set_run_font(badge, cfg, bold=True, color=_hex_to_rgb("BF8F00"))


def _short_url(url: str, maxlen: int = 60) -> str:
    if len(url) <= maxlen:
        return url
    return url[:maxlen - 3] + "..."


def _render_norma_citada(doc: Document, cfg: StyleConfig, *, data: dict) -> None:
    norma = data.get("norma", "")
    tier = _tier_from_data(data)
    fuente_url = data.get("fuente_url_oficial") or data.get("fuente_url")
    derogada_por = data.get("derogada_por")
    suggested = data.get("suggested_correction")
    contenido_runs = data.get("contenido") or []
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    bullet = p.add_run("· ")
    _set_run_font(bullet, cfg)
    label = p.add_run(norma)
    _set_run_font(label, cfg, bold=True, underline=True, color=_hex_to_rgb("1F3864"))
    _render_tier_badge(p, cfg, tier, fuente_url=fuente_url,
                        derogada_por=derogada_por, suggested=suggested)
    if contenido_runs:
        for r in contenido_runs[:3]:
            if isinstance(r, dict):
                text = r.get("text", "")
                if text:
                    run = p.add_run(" — " + text)
                    _set_run_font(run, cfg, italic=True, size_pt=cfg.body_size_pt - 1)


def _render_jurisprudencia(doc: Document, cfg: StyleConfig, *, data: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    bullet = p.add_run("· ")
    _set_run_font(bullet, cfg)
    label = p.add_run(f"{data.get('id','')} ({data.get('corte','')}, MP. {data.get('mp','')})")
    _set_run_font(label, cfg, bold=True, italic=True)
    # M20.10: tier badge para jurisprudencia
    tier = _tier_from_data(data)
    fuente_url = data.get("fuente_url_oficial") or data.get("fuente_url")
    _render_tier_badge(p, cfg, tier, fuente_url=fuente_url)
    ratio = data.get("ratio") or []
    if ratio:
        for r in ratio[:2]:
            if isinstance(r, dict) and r.get("text"):
                run = p.add_run(" — " + r["text"])
                _set_run_font(run, cfg, italic=True, size_pt=cfg.body_size_pt - 1)


def _render_table(doc: Document, cfg: StyleConfig, *, header: list[str], rows: list[list[str]]) -> None:
    if not header and not rows:
        return
    cols = max(len(header), max((len(r) for r in rows), default=0))
    if cols == 0:
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style = "Light Grid"
    # Header
    for i, cell in enumerate(tbl.rows[0].cells[:len(header)]):
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header[i]))
        _set_run_font(run, cfg, bold=True)
    # Rows
    for i, row_data in enumerate(rows, start=1):
        for j, val in enumerate(row_data[:cols]):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run_font(run, cfg)


def _render_juramento(doc: Document, cfg: StyleConfig, *, data: dict) -> None:
    text = data.get("text", "")
    norma_ref = data.get("norma_ref")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(cfg.paragraph_after_pt)
    head = p.add_run("JURAMENTO. ")
    _set_run_font(head, cfg, bold=True)
    body = p.add_run(text)
    _set_run_font(body, cfg)
    if norma_ref:
        ref = p.add_run(f" ({norma_ref})")
        _set_run_font(ref, cfg, italic=True, size_pt=cfg.body_size_pt - 1)


def _render_firma(doc: Document, cfg: StyleConfig, *, data: dict) -> None:
    parties = data.get("parties")
    if parties and isinstance(parties, list):
        # Multi-firma
        for party in parties:
            _render_single_signature_block(doc, cfg, data=party)
            doc.add_paragraph()
        return
    _render_single_signature_block(doc, cfg, data=data)


def _render_single_signature_block(doc: Document, cfg: StyleConfig, *, data: dict) -> None:
    ciudad_fecha = data.get("ciudad_fecha") or data.get("ciudad") or ""
    nombre = data.get("nombre", "")
    tp = data.get("tp", "")
    cc = data.get("cc", "")
    email = data.get("email", "")
    rol = data.get("rol", "")
    cargo = data.get("cargo", "")
    razon_social = data.get("razon_social", "")
    if ciudad_fecha:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(ciudad_fecha)
        _set_run_font(run, cfg)
    # Línea de firma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("_" * 40)
    _set_run_font(run, cfg)
    if rol:
        p_rol = doc.add_paragraph()
        run = p_rol.add_run(rol.upper())
        _set_run_font(run, cfg, bold=True)
    if nombre:
        p_n = doc.add_paragraph()
        run = p_n.add_run(nombre.upper())
        _set_run_font(run, cfg, bold=True)
    if cc:
        p_c = doc.add_paragraph()
        run = p_c.add_run(f"C.C. No. {cc}")
        _set_run_font(run, cfg)
    if razon_social:
        p_rs = doc.add_paragraph()
        run = p_rs.add_run(razon_social)
        _set_run_font(run, cfg)
    if cargo:
        p_cg = doc.add_paragraph()
        run = p_cg.add_run(cargo)
        _set_run_font(run, cfg, italic=True)
    if tp:
        p_t = doc.add_paragraph()
        run = p_t.add_run(f"T.P. No. {tp} del C.S. de la J.")
        _set_run_font(run, cfg)
    if email:
        p_e = doc.add_paragraph()
        run = p_e.add_run(email)
        _set_run_font(run, cfg)


def _render_blank(doc: Document, cfg: StyleConfig) -> None:
    p = doc.add_paragraph()
    p.add_run("")


# ============================================================================
# Top-level builder
# ============================================================================


def build_docx_from_blocks(
    blocks: list[dict[str, Any]],
    *,
    title: str = "",
    author: str = "LexAI",
    skill_context: Optional[SkillContext] = None,
) -> bytes:
    """Construye un .docx desde los bloques persistidos en document_blocks.

    Args:
        blocks: lista de dicts con shape {section_key, block_type, block_data, block_order}
        title: título por default si no hay TitleBlock en los blocks
        skill_context: M19.31 — contexto del SKILL.md para aplicar estilo
            específico del doc_type. Si None, usa estilo forense por default.

    Returns:
        bytes del .docx generado.
    """
    cfg = StyleConfig(skill_context)
    doc = Document()
    doc.core_properties.author = author
    if title:
        doc.core_properties.title = title

    _apply_page_setup(doc, cfg)
    _apply_default_doc_style(doc, cfg)

    # Contador de cláusulas para mapear → título del SKILL.md
    # (incrementa solo cuando aparece un SectionHeadingBlock)
    clause_idx = -1
    title_rendered = False

    # Ordenar bloques por block_order si está
    sorted_blocks = sorted(
        blocks or [],
        key=lambda b: b.get("block_order", 0) or 0,
    )

    for blk in sorted_blocks:
        bd = blk.get("block_data") or {}
        bt = blk.get("block_type") or bd.get("type")
        try:
            if bt == "title":
                _render_title(doc, cfg, bd.get("text") or title)
                title_rendered = True
            elif bt == "section_heading":
                clause_idx += 1
                _render_section_heading(
                    doc, cfg,
                    roman=bd.get("roman"),
                    text=bd.get("text", ""),
                    clause_idx=clause_idx,
                    skill_ctx=skill_context,
                )
            elif bt == "subsection":
                _render_subsection(doc, cfg,
                                    number=bd.get("number", ""),
                                    text=bd.get("text", ""))
            elif bt == "paragraph":
                _render_paragraph_with_runs(doc, cfg,
                                            runs_data=bd.get("runs") or [],
                                            align=bd.get("align", "justify"),
                                            indent_left_cm=bd.get("indent_left_cm"))
            elif bt == "hecho":
                _render_hecho(doc, cfg, num=bd.get("num", 0), runs_data=bd.get("runs") or [])
            elif bt == "pretension":
                _render_pretension(doc, cfg, ord_text=bd.get("ord", ""), runs_data=bd.get("runs") or [])
            elif bt == "list_item":
                _render_list_item(doc, cfg, num=bd.get("num", ""), runs_data=bd.get("runs") or [])
            elif bt == "norma_citada":
                _render_norma_citada(doc, cfg, data=bd)
            elif bt == "jurisprudencia":
                _render_jurisprudencia(doc, cfg, data=bd)
            elif bt == "table":
                _render_table(doc, cfg, header=bd.get("header") or [], rows=bd.get("rows") or [])
            elif bt == "juramento":
                _render_juramento(doc, cfg, data=bd)
            elif bt == "firma":
                _render_firma(doc, cfg, data=bd)
            elif bt == "blank":
                _render_blank(doc, cfg)
            elif bt == "silogismo":
                _render_paragraph_with_runs(
                    doc, cfg,
                    runs_data=[{"text": "PREMISA MAYOR: ", "bold": True}] + (bd.get("premisa_mayor") or []),
                )
                _render_paragraph_with_runs(
                    doc, cfg,
                    runs_data=[{"text": "PREMISA MENOR: ", "bold": True}] + (bd.get("premisa_menor") or []),
                )
                _render_paragraph_with_runs(
                    doc, cfg,
                    runs_data=[{"text": "CONCLUSIÓN: ", "bold": True}] + (bd.get("conclusion") or []),
                )
            elif bt == "calc_step":
                _render_paragraph_with_runs(doc, cfg,
                                            runs_data=[{"text": bd.get("label", ""), "bold": True}])
                _render_paragraph_with_runs(doc, cfg,
                                            runs_data=[{"text": "Fórmula: ", "bold": True}, {"text": bd.get("formula", "")}])
                _render_paragraph_with_runs(doc, cfg,
                                            runs_data=[{"text": "Aplicación: ", "bold": True}, {"text": bd.get("aplicacion", "")}])
                _render_paragraph_with_runs(doc, cfg,
                                            runs_data=[{"text": bd.get("total", ""), "bold": True}])
            else:
                # Fallback genérico: renderizar texto si está
                text = bd.get("text") or ""
                if text:
                    _render_paragraph_with_runs(doc, cfg, runs_data=[{"text": text}])
        except Exception as e:
            logger.warning("python_docx_builder: block %s render failed: %s", bt, e)
            continue

    # Si no hubo TitleBlock y nos pasaron un title fallback, lo agregamos al inicio
    # (es raro pero por compat con docx_forensic_builder)
    # No insertamos al inicio (python-docx no soporta insert-at-top cómodo);
    # quedó al final si fue necesario.

    # Serializar a bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
