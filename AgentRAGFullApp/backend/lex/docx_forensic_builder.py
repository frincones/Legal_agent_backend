"""DocX Forensic Builder — imperativo a partir de bloques tipados.

Replica el patrón de Claude (Anthropic) con docx.js — pero en Python con python-docx.

M19.10 mejoras:
  - Colores Heading 1 (#1F3864) y Heading 2 (#2F5496) estilo profesional
  - Hyperlinks azul/underline (#0563C1) en norma_citada y jurisprudencia (desde fuente_ref/fuente_url)
  - Shading en tablas: header azul oscuro + total ámbar + filas alternas grises
  - Footer con paginación dinámica "Página X de Y - Título"
  - Borders horizontales después de cada section_heading

M19.11 mejoras:
  - Numbering nativo Word (bullets y decimal) via docx.oxml.ns
  - Generador automático de tabla de liquidación desde calc_step blocks

WYSIWYG: mismo árbol de bloques que renderiza el canvas frontend.
"""
from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)

FONT = "Times New Roman"
# M19.14.C: alineado con referencia forense Colombia (demanda laboral)
# body 11pt, subheadings 12pt, headings 14pt, Letter size, márgenes 1 in (2.54 cm)
FONT_SIZE_PT = 11
FONT_SIZE_SUBHEADING = 12
FONT_SIZE_HEADING = 14

# Paleta profesional (estilo Claude/Word default)
COLOR_H1 = (0x1F, 0x38, 0x64)        # Azul oscuro corporate
COLOR_H2 = (0x2F, 0x54, 0x96)        # Azul medio
COLOR_LINK = (0x05, 0x63, 0xC1)      # Azul hyperlink Word default
COLOR_DEROGADA = (0xC0, 0x00, 0x00)  # Rojo derogada (preservar)
COLOR_TABLE_HEADER_BG = "1F3864"      # Header azul oscuro (sin #)
COLOR_TABLE_HEADER_FG = (0xFF, 0xFF, 0xFF)  # Texto blanco en header
COLOR_TABLE_TOTAL_BG = "D5E8F0"       # Total row ámbar/azul claro
COLOR_TABLE_ALT_BG = "F7F7F4"         # Filas alternas crema
COLOR_BORDER_SEP = "999999"           # Border separator gris


def build_docx_from_blocks(
    blocks: list[dict[str, Any]],
    title: str = "Documento",
    author: str = "LexAI",
) -> bytes:
    """Construye .docx forense a partir de lista de bloques tipados."""
    from docx import Document
    from docx.shared import Cm, Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author

    # M19.14.C: Letter size + márgenes 1 inch (estándar forense Colombia)
    # — alineado con referencia Demanda_Laboral_Gutierrez_vs_Constructora_Andina.docx
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(FONT_SIZE_PT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # M19.14.C: line_spacing 1.15 (más compacto, estilo forense moderno)
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    # M19.14.C: sin first_line_indent (la referencia no lo usa)
    pf.first_line_indent = Cm(0)

    # M19.10.A1 + M19.14.C: Headings alineados a referencia
    heading_styles = [
        (1, FONT_SIZE_HEADING, COLOR_H1),   # H1: 14pt #1F3864 azul oscuro
        (2, FONT_SIZE_SUBHEADING, COLOR_H2), # H2: 12pt #2F5496 azul medio
        (3, FONT_SIZE_SUBHEADING, COLOR_H2), # H3: mismo H2
    ]
    for lvl, size, color in heading_styles:
        try:
            h = doc.styles[f"Heading {lvl}"]
            h.font.name = FONT
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(*color)
        except Exception:
            pass

    # M19.11.B2: pre-procesar blocks para auto-generar tabla de liquidación
    # Si hay calc_step blocks consecutivos en la sección "liquidacion", agregar
    # tabla resumen al final con totales.
    blocks = _augment_with_liquidacion_table(blocks)

    # M19.10.A7: registry de derogation links (norma_vigente_url por block_id)
    # Permite que norma_citada muestre badge derogada + link a la vigente
    derog_lookup = _build_derogation_lookup(blocks)

    # M19.10.A5: Footer con paginación
    _add_footer_with_pagination(doc, title)

    # Render bloques
    for b in blocks:
        bt = b.get("block_type") or b.get("block_data", {}).get("type")
        bd = b.get("block_data") or {}
        try:
            _render_block(doc, bt, bd, derog_lookup)
        except Exception as e:
            logger.warning("render block %s failed: %s", bt, e)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_derogation_lookup(blocks: list[dict[str, Any]]) -> dict[str, str]:
    """Mapa block_id -> fuente_url_vigente (de norma que reemplaza la derogada)."""
    lookup: dict[str, str] = {}
    for b in blocks:
        bd = b.get("block_data") or {}
        if bd.get("type") == "norma_citada" and bd.get("derogada") and bd.get("fuente_url_vigente"):
            block_id = b.get("block_id") or bd.get("block_id")
            if block_id:
                lookup[block_id] = bd["fuente_url_vigente"]
    return lookup


def _augment_with_liquidacion_table(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """M19.11.B2: si hay ≥3 calc_step en sección 'liquidacion' o 'estimacion_cuantia_liquidacion',
    insertar tabla resumen automática justo después del último calc_step de esa sección.

    Idempotente: si ya existe table block después de los calc_step, no duplica.
    """
    LIQ_SECTIONS = {"liquidacion", "estimacion_cuantia_liquidacion", "estimacion_cuantia_y_liquidacion"}
    # Agrupar calc_steps por section_key
    calc_by_section: dict[str, list[dict]] = {}
    for b in blocks:
        if (b.get("block_type") or b.get("block_data", {}).get("type")) == "calc_step":
            sk = b.get("section_key") or "default"
            calc_by_section.setdefault(sk, []).append(b)

    # Para cada sección de liquidación con ≥3 calc_steps, generar tabla resumen
    augmented = list(blocks)
    for sk, calcs in calc_by_section.items():
        if sk not in LIQ_SECTIONS or len(calcs) < 3:
            continue

        # Verificar si ya hay table block después del último calc_step en esa sección
        last_idx = augmented.index(calcs[-1])
        next_block = augmented[last_idx + 1] if last_idx + 1 < len(augmented) else None
        if next_block and (next_block.get("block_type") or next_block.get("block_data", {}).get("type")) == "table":
            continue  # ya hay tabla resumen, skip

        # Construir tabla resumen
        rows = []
        total_estimado = 0.0
        for c in calcs:
            cd = c.get("block_data") or {}
            label = cd.get("label", "")
            formula = cd.get("formula", "")
            total_text = cd.get("total", "")
            rows.append([label, formula, total_text])
            # Intentar parsear monto del total para sumar
            import re
            nums = re.findall(r"[\d.,]+", total_text.replace("$", ""))
            if nums:
                try:
                    val = float(nums[0].replace(".", "").replace(",", "."))
                    total_estimado += val
                except Exception:
                    pass

        if total_estimado > 0:
            rows.append([
                "TOTAL ESTIMADO",
                "Suma de conceptos liquidados",
                f"${total_estimado:,.0f}".replace(",", "."),
            ])

        table_block = {
            "section_key": sk,
            "block_order": (calcs[-1].get("block_order") or 0) + 1,
            "block_id": f"auto_liq_table_{sk}",
            "block_type": "table",
            "block_data": {
                "type": "table",
                "header": ["CONCEPTO", "BASE / FÓRMULA", "VALOR APROXIMADO (COP)"],
                "rows": rows,
                "has_total_row": total_estimado > 0,
                "header_shading": COLOR_TABLE_HEADER_BG,
                "total_row_shading": COLOR_TABLE_TOTAL_BG,
                "alternate_row_shading": COLOR_TABLE_ALT_BG,
                "auto_generated": True,
            },
        }
        augmented.insert(last_idx + 1, table_block)

    return augmented


def _add_footer_with_pagination(doc, title: str) -> None:
    """M19.10.A5: Footer con 'TITULO - Página X de Y'."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clear existing
        for elem in p._element.findall(qn("w:r")):
            p._element.remove(elem)

        # Texto: título corto
        truncated = (title[:60] + "…") if len(title) > 60 else title
        r1 = p.add_run(f"{truncated} — Página ")
        r1.font.name = FONT
        r1.font.size = Pt(9)
        r1.italic = True

        # Campo PAGE
        _add_field(p, "PAGE")

        r2 = p.add_run(" de ")
        r2.font.name = FONT
        r2.font.size = Pt(9)
        r2.italic = True

        _add_field(p, "NUMPAGES")


def _add_field(paragraph, field_code: str) -> None:
    """Inserta un campo dinámico Word (ej. PAGE, NUMPAGES)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(9)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f" {field_code} "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False) -> None:
    """M19.10.A2: Inserta hyperlink azul/underline en el párrafo.

    Genera el XML correcto de python-docx con w:hyperlink.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    if not url:
        # fallback: add as regular bold/italic run
        run = paragraph.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(FONT_SIZE_PT)
        if bold: run.bold = True
        if italic: run.italic = True
        return

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Estilo Hyperlink (azul + underline)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:i")
        rPr.append(i)

    # Font + size
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(FONT_SIZE_PT * 2))  # half-points
    rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_cell_shading(cell, fill_color_hex: str) -> None:
    """M19.10.A4: Aplica color de fondo a una celda. fill='1F3864' (sin #)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill_color_hex)
    tc_pr.append(shading)


def _add_horizontal_border(paragraph) -> None:
    """M19.10.A6: Border bottom horizontal en un párrafo (separator visual)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_BORDER_SEP)
    pBdr.append(bottom)

    pPr.append(pBdr)


def _add_native_numbering(paragraph, num_format: str = "decimal", level: int = 0) -> None:
    """M19.11.B1: Aplica numbering nativo Word a un párrafo.

    num_format: 'decimal' (1./2./3.) o 'bullet' (•)
    NOTA: para listas nativas con numId, python-docx requiere definir numId
    en numbering.xml — sin custom, usamos los styles predefinidos 'List Number'
    / 'List Bullet' que ya están en cualquier doc nuevo.
    """
    try:
        style_name = "List Bullet" if num_format == "bullet" else "List Number"
        paragraph.style = paragraph.part.document.styles[style_name]
    except Exception:
        # Si el style no existe, fallback silencioso (numbering manual del block sigue funcionando)
        pass


def _render_block(doc, btype: str, bd: dict[str, Any], derog_lookup: dict[str, str] | None = None) -> None:
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    if btype == "title":
        text = (bd.get("text") or "").upper()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(*COLOR_H1)  # M19.10.A1: título azul
        run.font.size = Pt(16) if bd.get("level", 1) <= 1 else Pt(14)
        return

    if btype == "section_heading":
        text = f"{bd.get('roman', '')}. {(bd.get('text') or '').upper()}"
        h = doc.add_heading(text, level=1)  # M19.10.A1: usar H1 con color
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        # M19.10.A6: border separator después del heading
        sep = doc.add_paragraph()
        sep.paragraph_format.first_line_indent = Cm(0)
        sep.paragraph_format.space_after = Pt(8)
        _add_horizontal_border(sep)
        return

    if btype == "subsection":
        text = f"{bd.get('number', '')}. {bd.get('text', '')}"
        h = doc.add_heading(text, level=2)
        return

    if btype == "paragraph":
        p = doc.add_paragraph()
        align = bd.get("align", "justify")
        p.alignment = {
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
        }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_runs(p, bd.get("runs", []))
        return

    if btype == "hecho":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1)
        # M19.11.B1: numbering nativo si block tiene flag use_native_numbering
        # (default mantiene texto manual para retrocompat)
        if bd.get("use_native_numbering"):
            _add_native_numbering(p, "decimal")
        else:
            run = p.add_run(f"{bd.get('num', '?')}.\t")
            run.bold = True
            run.font.name = FONT
            run.font.size = Pt(FONT_SIZE_PT)
        _add_runs(p, bd.get("runs", []))
        return

    if btype == "pretension":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"{bd.get('ord', '?')}.- ")
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(FONT_SIZE_PT)
        _add_runs(p, bd.get("runs", []))
        return

    if btype == "norma_citada":
        # M19.10.A2: si fuente_ref/fuente_url, convertir a hyperlink
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run("  · ")
        norma_text = bd.get("norma", "")
        url = bd.get("fuente_ref") or bd.get("fuente_url") or bd.get("fuente_url_original")

        if url:
            _add_hyperlink(p, norma_text, url, bold=True)
        else:
            r = p.add_run(norma_text)
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(FONT_SIZE_PT)

        if bd.get("derogada"):
            r2 = p.add_run(" [DEROGADA")
            r2.font.color.rgb = RGBColor(*COLOR_DEROGADA)
            r2.italic = True
            r2.bold = True
            # Si tenemos URL vigente, agregar link "ver vigente"
            block_id = bd.get("block_id", "")
            url_vigente = (derog_lookup or {}).get(block_id) or bd.get("fuente_url_vigente")
            if url_vigente:
                r3 = p.add_run(" → ")
                r3.italic = True
                _add_hyperlink(p, "norma vigente", url_vigente, italic=True)
            r4 = p.add_run("]")
            r4.font.color.rgb = RGBColor(*COLOR_DEROGADA)
            r4.italic = True
            r4.bold = True
        elif bd.get("verified"):
            r2 = p.add_run(" ✓")
            r2.bold = True
            r2.font.color.rgb = RGBColor(0x0D, 0x7A, 0x3D)  # verde verified
        return

    if btype == "jurisprudencia":
        # M19.10.A3: hyperlink en sentencia desde fuente_url
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run("  · ")
        sentencia_id = bd.get("id", "")
        mp = bd.get("mp", "")
        corte = bd.get("corte", "")
        url = bd.get("fuente_url") or bd.get("fuente_ref")

        if url:
            _add_hyperlink(p, sentencia_id, url, bold=True)
        else:
            r = p.add_run(sentencia_id)
            r.bold = True
            r.font.name = FONT
            r.font.size = Pt(FONT_SIZE_PT)
        r2 = p.add_run(f", M.P. {mp} ({corte})")
        r2.bold = True
        r2.font.name = FONT
        r2.font.size = Pt(FONT_SIZE_PT)

        if bd.get("ratio"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(1.0)
            p2.paragraph_format.first_line_indent = Cm(0)
            r3 = p2.add_run("    “")
            r3.italic = True
            _add_runs(p2, bd.get("ratio", []), italic_default=True)
            r4 = p2.add_run("”")
            r4.italic = True
        return

    if btype == "silogismo":
        for label, runs_key in [
            ("Premisa Mayor (norma): ", "premisa_mayor"),
            ("Premisa Menor (hecho): ", "premisa_menor"),
            ("Conclusión: ", "conclusion"),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(label)
            r.bold = True
            _add_runs(p, bd.get(runs_key, []))
        return

    if btype == "table":
        # M19.10.A4: tabla con shading profesional (header azul, total ámbar, alternas)
        header = bd.get("header", [])
        rows = bd.get("rows", [])
        if not header or not rows:
            return

        tbl = doc.add_table(rows=len(rows) + 1, cols=len(header))
        tbl.style = "Table Grid"  # bordes finos visibles

        header_shading = bd.get("header_shading") or COLOR_TABLE_HEADER_BG
        total_shading = bd.get("total_row_shading") or COLOR_TABLE_TOTAL_BG
        alt_shading = bd.get("alternate_row_shading") or COLOR_TABLE_ALT_BG
        has_total = bd.get("has_total_row", False)

        # Header row con shading azul + texto blanco
        for ci, h in enumerate(header):
            cell = tbl.rows[0].cells[ci]
            _set_cell_shading(cell, header_shading)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            r = p.add_run(h)
            r.font.name = FONT
            r.font.size = Pt(11)
            r.bold = True
            r.font.color.rgb = RGBColor(*COLOR_TABLE_HEADER_FG)  # blanco
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for ri, row in enumerate(rows):
            is_total = has_total and ri == len(rows) - 1
            # Alternate shading para filas no-total
            if is_total:
                shade = total_shading
            elif ri % 2 == 1:
                shade = alt_shading
            else:
                shade = None

            for ci in range(len(header)):
                cell = tbl.rows[ri + 1].cells[ci]
                if shade:
                    _set_cell_shading(cell, shade)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = None
                # Alinear última columna a la derecha (típicamente "VALOR")
                if ci == len(header) - 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                txt = row[ci] if ci < len(row) else ""
                r = p.add_run(str(txt))
                r.font.name = FONT
                r.font.size = Pt(11)
                if is_total:
                    r.bold = True
        return

    if btype == "calc_step":
        p1 = doc.add_paragraph()
        p1.paragraph_format.first_line_indent = Cm(0)
        r = p1.add_run(bd.get("label", ""))
        r.bold = True
        r.underline = True

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        r2 = p2.add_run(f"Fórmula: {bd.get('formula', '')}")
        r2.italic = True

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.first_line_indent = Cm(0)
        r3 = p3.add_run(f"Aplicación: {bd.get('aplicacion', '')}")
        r3.italic = True

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.paragraph_format.first_line_indent = Cm(0)
        r4 = p4.add_run(bd.get("total", ""))
        r4.bold = True
        return

    if btype == "list_item":
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.5)
        # M19.11.B1: numbering nativo si flag presente
        if bd.get("use_native_numbering") == "bullet":
            _add_native_numbering(p, "bullet")
            _add_runs(p, bd.get("runs", []))
        elif bd.get("use_native_numbering"):
            _add_native_numbering(p, "decimal")
            _add_runs(p, bd.get("runs", []))
        else:
            r = p.add_run(f"{bd.get('num', '')}) ")
            r.bold = True
            _add_runs(p, bd.get("runs", []))
        return

    if btype == "juramento":
        doc.add_paragraph("")
        h = doc.add_heading("JURAMENTO", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Border separator después del heading
        sep = doc.add_paragraph()
        sep.paragraph_format.first_line_indent = Cm(0)
        sep.paragraph_format.space_after = Pt(8)
        _add_horizontal_border(sep)

        if bd.get("norma_ref"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(bd["norma_ref"])
            r.italic = True
        p2 = doc.add_paragraph(bd.get("text", ""))
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return

    if btype == "firma":
        doc.add_paragraph("")
        doc.add_paragraph("")
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.first_line_indent = Cm(0)
        r = p1.add_run(bd.get("ciudad_fecha", ""))
        r.italic = True

        doc.add_paragraph("")
        doc.add_paragraph("Atentamente,")
        doc.add_paragraph("")
        doc.add_paragraph("_____________________________________")

        p2 = doc.add_paragraph()
        p2.paragraph_format.first_line_indent = Cm(0)
        r2 = p2.add_run(bd.get("nombre", ""))
        r2.bold = True

        p3 = doc.add_paragraph()
        p3.paragraph_format.first_line_indent = Cm(0)
        p3.add_run(f"Abogado · T.P. No. {bd.get('tp', '')} del C.S.J.")

        if bd.get("cc"):
            p4 = doc.add_paragraph()
            p4.paragraph_format.first_line_indent = Cm(0)
            p4.add_run(f"C.C. No. {bd.get('cc')}")
        contact = []
        if bd.get("email"):
            contact.append(f"Email: {bd['email']}")
        if bd.get("telefono"):
            contact.append(f"Tel.: {bd['telefono']}")
        if contact:
            p5 = doc.add_paragraph()
            p5.paragraph_format.first_line_indent = Cm(0)
            p5.add_run(" · ".join(contact))
        return

    if btype == "blank":
        doc.add_paragraph("")
        return


def _add_runs(paragraph, runs: list[dict[str, Any]], italic_default: bool = False) -> None:
    """Inserta runs con bold/italic/underline. Detecta urls inline `[texto](url)`."""
    from docx.shared import Pt
    import re

    URL_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")

    for r in runs or []:
        if not isinstance(r, dict):
            continue
        text = r.get("text", "")
        if not text:
            continue
        url = r.get("url") or r.get("fuente_url")
        bold = bool(r.get("bold"))
        italic = bool(r.get("italic") or italic_default)
        underline = bool(r.get("underline"))

        # Si run trae url explícito, hyperlink
        if url:
            _add_hyperlink(paragraph, text, url, bold=bold, italic=italic)
            continue

        # M19.10: detectar markdown links inline `[texto](url)`
        last_pos = 0
        matches = list(URL_PATTERN.finditer(text))
        if matches:
            for m in matches:
                # texto antes del link
                if m.start() > last_pos:
                    plain = text[last_pos:m.start()]
                    if plain:
                        run = paragraph.add_run(plain)
                        run.font.name = FONT
                        run.font.size = Pt(FONT_SIZE_PT)
                        if bold: run.bold = True
                        if italic: run.italic = True
                        if underline: run.underline = True
                # hyperlink
                _add_hyperlink(paragraph, m.group(1), m.group(2), bold=bold, italic=italic)
                last_pos = m.end()
            # tail después del último link
            if last_pos < len(text):
                tail = text[last_pos:]
                if tail:
                    run = paragraph.add_run(tail)
                    run.font.name = FONT
                    run.font.size = Pt(FONT_SIZE_PT)
                    if bold: run.bold = True
                    if italic: run.italic = True
                    if underline: run.underline = True
            continue

        # Run regular
        run = paragraph.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(FONT_SIZE_PT)
        if bold: run.bold = True
        if italic: run.italic = True
        if underline: run.underline = True
