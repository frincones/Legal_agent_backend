"""Sprint E · Router /v1/canvas/export · genera .docx desde markdown.

Usa python-docx para construir el documento. Soporta:
  - Plain markdown → .docx con styles
  - Markdown + tracked changes (redlines aplicados como revisions)
  - style="forensic" para formato bufete top de Bogotá:
      Times New Roman 12pt, justificado, márgenes 3cm,
      headings centrados en MAYÚSCULAS bold, tablas, footer firma+TP.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from utils.auth import Principal, get_current_firm

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/v1/canvas", tags=["canvas-export"])


class ExportBody(BaseModel):
    content_md: str = Field(..., min_length=1)
    title: Optional[str] = "Documento"
    author: Optional[str] = "LexAI"
    redlines_applied: Optional[bool] = False
    format: str = Field(default="docx", pattern="^(docx|pdf|md)$")
    style: str = Field(default="forensic", pattern="^(forensic|plain)$")
    # Datos para el bloque de firma final (opcionales)
    abogado_nombre: Optional[str] = None
    abogado_tp: Optional[str] = None
    abogado_cc: Optional[str] = None
    abogado_email: Optional[str] = None
    abogado_telefono: Optional[str] = None
    ciudad_fecha: Optional[str] = None


@router.post("/export")
async def export_canvas(
    body: ExportBody,
    principal: Principal = Depends(get_current_firm),
):
    fmt = body.format
    if fmt == "md":
        buf = io.BytesIO(body.content_md.encode("utf-8"))
        return StreamingResponse(
            buf,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=\"{_safe(body.title)}.md\""},
        )
    if fmt == "docx":
        try:
            buf = _build_docx(
                body.content_md,
                title=body.title or "Documento",
                author=body.author or "LexAI",
                style=body.style or "forensic",
                signature={
                    "nombre": body.abogado_nombre,
                    "tp": body.abogado_tp,
                    "cc": body.abogado_cc,
                    "email": body.abogado_email,
                    "telefono": body.abogado_telefono,
                    "ciudad_fecha": body.ciudad_fecha,
                },
            )
        except ImportError:
            raise HTTPException(500, "python-docx no instalado · pip install python-docx")
        except Exception as e:
            logger.exception("export docx failed")
            raise HTTPException(500, f"Error generando .docx: {str(e)[:120]}")
        return StreamingResponse(
            io.BytesIO(buf),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{_safe(body.title)}.docx\""},
        )
    if fmt == "pdf":
        raise HTTPException(501, "Export PDF próximamente · usa .docx por ahora")
    raise HTTPException(400, "Formato no soportado")


def _safe(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", name)[:80]


def _build_docx(
    md_text: str,
    title: str,
    author: str,
    style: str = "forensic",
    signature: dict | None = None,
) -> bytes:
    """Convierte markdown a .docx.

    style="forensic" produce formato bufete top:
      - Times New Roman 12pt
      - Márgenes 3cm sup/inf, 3cm izq, 2.5cm der (estándar forense Col.)
      - Body justificado, interlineado 1.5
      - H1 centrado MAYÚSCULAS bold (para "DEMANDA ORDINARIA LABORAL")
      - H2 izquierda numeración romana ("I. PARTES DEL PROCESO")
      - H3 subsecciones bold
      - Tablas markdown (| col | col |) renderizadas como tablas Word
      - Bloque de firma final con TP / CC / email del apoderado
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author

    forensic = style == "forensic"

    if forensic:
        # Page setup forense Colombia
        for section in doc.sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

        # Default Normal style → Times New Roman 12pt justified 1.5
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        pf = normal.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.5
        pf.space_after = Pt(6)
        pf.first_line_indent = Cm(1.25)

        # Heading styles
        for lvl, size in [(1, 14), (2, 13), (3, 12)]:
            try:
                h = doc.styles[f"Heading {lvl}"]
                h.font.name = "Times New Roman"
                h.font.size = Pt(size)
                h.font.bold = True
                h.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            except Exception:
                pass

    # Encabezado (título principal centrado MAYÚSCULAS si forensic)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run((title or "").upper() if forensic else title)
    title_run.bold = True
    title_run.font.size = Pt(16) if forensic else Pt(14)
    if forensic:
        title_p.paragraph_format.space_after = Pt(18)
        title_p.paragraph_format.first_line_indent = Cm(0)

    lines = md_text.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            i += 1
            continue

        # Tablas markdown (| h1 | h2 |\n|---|---|\n| a | b |)
        if forensic and line.startswith("|") and line.endswith("|") and (i + 1 < n) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].rstrip()):
            table_lines = [line]
            i += 2  # skip separator
            while i < n and lines[i].rstrip().startswith("|") and lines[i].rstrip().endswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        # Headers
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip().upper() if forensic else line[2:].strip(), level=1)
            if forensic:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        if line.startswith("## "):
            text = line[3:].strip()
            h = doc.add_heading(text.upper() if forensic else text, level=2)
            if forensic:
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Bullets
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
            if forensic:
                p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, re.sub(r"^\d+\.\s", "", line))
            if forensic:
                p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        # Paragraph con inline emphasis
        p = doc.add_paragraph()
        _add_runs(p, line)
        i += 1

    # Bloque de firma forense
    if forensic and signature and any(signature.values()):
        doc.add_paragraph("")
        doc.add_paragraph("")
        ciudad = signature.get("ciudad_fecha") or "Bogotá D.C., a la fecha de presentación."
        p_ciudad = doc.add_paragraph()
        p_ciudad.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ciudad.paragraph_format.first_line_indent = Cm(0)
        p_ciudad.add_run(ciudad).italic = True

        doc.add_paragraph("")
        doc.add_paragraph("")
        p_resp = doc.add_paragraph()
        p_resp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_resp.paragraph_format.first_line_indent = Cm(0)
        p_resp.add_run("Atentamente,").bold = False

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("_____________________________________")

        nombre = signature.get("nombre") or "[NOMBRE_APODERADO]"
        tp = signature.get("tp") or "[TP_APODERADO]"
        cc = signature.get("cc")
        email = signature.get("email")
        tel = signature.get("telefono")

        p_nombre = doc.add_paragraph()
        p_nombre.paragraph_format.first_line_indent = Cm(0)
        p_nombre.add_run(nombre).bold = True

        p_tp = doc.add_paragraph()
        p_tp.paragraph_format.first_line_indent = Cm(0)
        p_tp.add_run(f"Abogado · T.P. No. {tp} del C.S.J.")

        if cc:
            p_cc = doc.add_paragraph()
            p_cc.paragraph_format.first_line_indent = Cm(0)
            p_cc.add_run(f"C.C. No. {cc}")
        if email or tel:
            contact_parts = []
            if email:
                contact_parts.append(f"Email: {email}")
            if tel:
                contact_parts.append(f"Tel.: {tel}")
            p_contact = doc.add_paragraph()
            p_contact.paragraph_format.first_line_indent = Cm(0)
            p_contact.add_run(" · ".join(contact_parts))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_markdown_table(doc, table_lines: list[str]) -> None:
    """Renderiza tabla markdown como tabla Word con bordes."""
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows = []
    for raw in table_lines:
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Light Grid Accent 1"
    for r_idx, row_data in enumerate(rows):
        for c_idx in range(ncols):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            txt = row_data[c_idx] if c_idx < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = None
            run = p.add_run(txt)
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            if r_idx == 0:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_runs(paragraph, text: str):
    """Inline parser bold (**...**) e italic (*...*) simple."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
